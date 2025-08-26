#!/usr/bin/env python3
"""
남은 BC카드(카드이용안내).docx 파일 처리
"""

import os
from docx import Document

def fix_remaining_file():
    """BC카드(카드이용안내).docx 파일에 구분자 추가"""
    s3_folder = "/mnt/d/99_DEOTIS_QA_SYSTEM/03_DEOTIS_QA/s3-chunking"
    filename = "BC카드(카드이용안내).docx"
    file_path = os.path.join(s3_folder, filename)
    
    print(f"📄 처리 중: {filename}")
    
    try:
        # 원본 파일 읽기
        doc = Document(file_path)
        
        # 새 문서 생성
        new_doc = Document()
        
        # 섹션 키워드 정의
        section_keywords = [
            "제1장", "제2장", "제3장", "제4장", "제5장",
            "1. 개요", "2. 신용카드", "3. 카드", "4. 결제", "5. 부가",
            "■", "▶", "○", "●"
        ]
        
        paragraphs = []
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                paragraphs.append(text)
        
        print(f"  └─ 총 {len(paragraphs)}개 문단 발견")
        
        # 첫 번째 전략: 제목 기반 섹션 분할
        current_section = []
        section_count = 0
        
        for i, text in enumerate(paragraphs):
            # 제목/섹션 시작 감지
            is_section_start = False
            
            # 키워드 매칭
            for keyword in section_keywords:
                if keyword in text and len(text) < 50:
                    is_section_start = True
                    break
            
            # 또는 번호가 있는 제목 (1., 1-1, 가. 등)
            if not is_section_start:
                import re
                if re.match(r'^[0-9]+\.', text) or re.match(r'^[0-9]+-[0-9]+', text) or re.match(r'^[가-힣]\.', text):
                    if len(text) < 50:
                        is_section_start = True
            
            if is_section_start and current_section:
                # 이전 섹션 저장
                for para in current_section:
                    new_doc.add_paragraph(para)
                
                # 구분자 추가 (첫 번째가 아닌 경우)
                if section_count > 0:
                    new_doc.add_paragraph("/$$/")
                    print(f"    ├─ 구분자 추가: 섹션 {section_count}")
                
                section_count += 1
                current_section = []
            
            current_section.append(text)
        
        # 마지막 섹션 처리
        if current_section:
            for para in current_section:
                new_doc.add_paragraph(para)
            section_count += 1
        
        # 구분자가 충분히 안 들어갔으면 강제로 추가
        if section_count < 5:
            print(f"  ⚠️ 섹션이 {section_count}개로 적음. 페이지 기반으로 재분할")
            
            # 새로운 전략: 문단 수 기반 분할
            new_doc = Document()
            chunk_size = len(paragraphs) // 8  # 약 8개 청크로 나누기
            
            for i, text in enumerate(paragraphs):
                new_doc.add_paragraph(text)
                
                # 일정 간격마다 구분자 추가
                if (i + 1) % chunk_size == 0 and i < len(paragraphs) - 1:
                    new_doc.add_paragraph("/$$/")
                    print(f"    ├─ 구분자 추가: {i//chunk_size + 1}번째 청크")
        
        # 파일 저장
        new_doc.save(file_path)
        print(f"✅ 파일 업데이트 완료: {filename}")
        print(f"   └─ 최종 섹션/청크 수: {section_count if section_count >= 5 else 8}")
        
        return True
        
    except Exception as e:
        print(f"❌ 처리 중 오류: {str(e)}")
        return False

if __name__ == "__main__":
    fix_remaining_file()