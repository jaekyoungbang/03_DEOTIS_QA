#!/usr/bin/env python3
"""
백업에서 원본 복원 후 구분자 추가
"""

import os
import shutil
from docx import Document

def final_fix():
    """백업에서 복원 후 구분자 추가"""
    s3_folder = "/mnt/d/99_DEOTIS_QA_SYSTEM/03_DEOTIS_QA/s3-chunking"
    backup_folder = os.path.join(s3_folder, "backup")
    
    # BC카드(카드이용안내).docx 처리
    filename = "BC카드(카드이용안내).docx"
    original_backup = os.path.join(backup_folder, f"original_{filename}")
    target_file = os.path.join(s3_folder, filename)
    
    print(f"📄 {filename} 최종 처리")
    
    try:
        # 1. 백업에서 원본 복원
        shutil.copy2(original_backup, target_file)
        print("  ✅ 원본 파일 복원됨")
        
        # 2. 문서 읽기
        doc = Document(target_file)
        
        # 3. 텍스트 추출
        all_text = []
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                all_text.append(text)
        
        print(f"  └─ 총 {len(all_text)}개 문단")
        
        # 4. 새 문서에 구분자를 적절히 배치하여 저장
        new_doc = Document()
        
        # 전체를 8개 청크로 나누기
        chunk_size = len(all_text) // 8
        
        for i, text in enumerate(all_text):
            new_doc.add_paragraph(text)
            
            # 매 chunk_size 마다 구분자 추가 (마지막 제외)
            if (i + 1) % chunk_size == 0 and i < len(all_text) - chunk_size:
                new_doc.add_paragraph("/$$/")
                print(f"    ├─ 구분자 추가: {(i+1)//chunk_size}번째 위치")
        
        # 5. 임시 파일로 저장
        temp_file = os.path.join(s3_folder, f"temp_{filename}")
        new_doc.save(temp_file)
        
        # 6. 원본을 임시 파일로 교체
        os.remove(target_file)
        os.rename(temp_file, target_file)
        
        print(f"✅ {filename} 구분자 추가 완료 (8개 청크)")
        
        return True
        
    except Exception as e:
        print(f"❌ 오류: {str(e)}")
        return False

if __name__ == "__main__":
    final_fix()