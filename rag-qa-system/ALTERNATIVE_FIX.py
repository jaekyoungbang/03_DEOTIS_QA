#!/usr/bin/env python3
"""
대안 해결책: docx2txt 없이 python-docx로 DOCX 파일 처리
"""

import os
import sys
from pathlib import Path

def fix_docx_loader():
    """document_processor.py 수정하여 python-docx 사용"""
    
    processor_file = Path("services/document_processor.py")
    
    if not processor_file.exists():
        print("❌ document_processor.py 파일을 찾을 수 없습니다.")
        return False
    
    # 백업 생성
    backup_file = processor_file.with_suffix('.py.backup')
    if not backup_file.exists():
        processor_file.rename(backup_file)
        print(f"✅ 백업 생성: {backup_file}")
    
    # 수정된 코드 작성
    new_content = '''from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import (
    TextLoader,
    UnstructuredMarkdownLoader
)
try:
    from langchain_community.document_loaders import PyPDFLoader
except ImportError:
    PyPDFLoader = None

# docx2txt 대신 python-docx 직접 사용
try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from langchain.schema import Document
from config import Config
import os
from typing import List

class DocumentProcessor:
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=Config.CHUNK_SIZE,
            chunk_overlap=Config.CHUNK_OVERLAP,
            length_function=len,
            separators=["\\n\\n", "\\n", " ", ""]
        )
        self.loaders = {
            '.txt': TextLoader,
            '.md': UnstructuredMarkdownLoader
        }
        if PyPDFLoader:
            self.loaders['.pdf'] = PyPDFLoader
        if DOCX_AVAILABLE:
            self.loaders['.docx'] = self._load_docx_custom
    
    def _load_docx_custom(self, file_path: str) -> List[Document]:
        """python-docx를 사용한 DOCX 파일 로더"""
        try:
            doc = docx.Document(file_path)
            full_text = []
            
            # 문단 텍스트 추출
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text)
            
            # 표 텍스트 추출
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        full_text.append(" | ".join(row_text))
            
            content = "\\n".join(full_text)
            
            if content.strip():
                metadata = {
                    'source': file_path,
                    'file_type': 'docx'
                }
                return [Document(page_content=content, metadata=metadata)]
            else:
                return []
                
        except Exception as e:
            raise Exception(f"Error loading DOCX document: {str(e)}")
    
    def load_document(self, file_path: str) -> List[Document]:
        """Load a document based on its file extension"""
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension not in self.loaders:
            raise ValueError(f"Unsupported file type: {file_extension}")
        
        loader_class_or_method = self.loaders[file_extension]
        
        try:
            if file_extension == '.docx':
                # 커스텀 DOCX 로더 사용
                documents = loader_class_or_method(file_path)
            else:
                # 기존 로더 사용
                loader = loader_class_or_method(file_path)
                documents = loader.load()
            return documents
        except Exception as e:
            raise Exception(f"Error loading document: {str(e)}")
    
    def split_documents(self, documents: List[Document]) -> List[Document]:
        """Split documents into chunks"""
        return self.text_splitter.split_documents(documents)
    
    def process_file(self, file_path: str, metadata: dict = None) -> List[Document]:
        """Process a file: load and split into chunks"""
        # Load the document
        documents = self.load_document(file_path)
        
        # Add custom metadata if provided
        if metadata:
            for doc in documents:
                doc.metadata.update(metadata)
        
        # Split into chunks
        chunks = self.split_documents(documents)
        
        # Add chunk index to metadata
        for i, chunk in enumerate(chunks):
            chunk.metadata['chunk_index'] = i
            chunk.metadata['source_file'] = os.path.basename(file_path)
        
        return chunks
    
    def process_text(self, text: str, metadata: dict = None) -> List[Document]:
        """Process raw text into chunks"""
        # Create a document from text
        doc = Document(page_content=text, metadata=metadata or {})
        
        # Split into chunks
        chunks = self.text_splitter.split_documents([doc])
        
        # Add chunk index to metadata
        for i, chunk in enumerate(chunks):
            chunk.metadata['chunk_index'] = i
        
        return chunks
    
    def validate_file(self, file_path: str) -> bool:
        """Validate if a file can be processed"""
        # Check if file exists
        if not os.path.exists(file_path):
            return False
        
        # Check file size
        file_size = os.path.getsize(file_path)
        if file_size > Config.MAX_FILE_SIZE:
            return False
        
        # Check file extension
        file_extension = os.path.splitext(file_path)[1].lower()
        if file_extension not in self.loaders:
            return False
        
        return True
'''
    
    # 새 파일 작성
    with open(processor_file, 'w', encoding='utf-8') as f:
        f.write(new_content)
    
    print("✅ document_processor.py가 수정되었습니다.")
    print("✅ 이제 python-docx를 사용하여 DOCX 파일을 처리합니다.")
    return True

if __name__ == "__main__":
    print("========================================")
    print("DOCX 로더 대안 수정")
    print("========================================")
    
    if fix_docx_loader():
        print("\\n✅ 수정 완료! 서버를 재시작하세요.")
        print("python app.py")
    else:
        print("\\n❌ 수정 실패")
    
    input("\\n계속하려면 Enter를 누르세요...")