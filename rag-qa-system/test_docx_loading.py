#!/usr/bin/env python3
"""
DOCX 로딩 테스트 및 문제 해결
"""

import os
import sys
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

def test_docx_import():
    """DOCX 관련 모듈 import 테스트"""
    print("=== DOCX 모듈 import 테스트 ===")
    
    try:
        import docx
        print("✅ python-docx 모듈 import 성공")
        
        # 테스트 문서 경로
        test_file = r"D:\99_DEOTIS_QA_SYSTEM\03_DEOTIS_QA\s3\BC카드(신용카드 업무처리 안내).docx"
        
        if os.path.exists(test_file):
            print(f"✅ 테스트 파일 존재: {test_file}")
            
            # 간단한 로딩 테스트
            doc = docx.Document(test_file)
            print(f"✅ DOCX 파일 로딩 성공 - 문단 수: {len(doc.paragraphs)}")
            
            # 첫 번째 문단 미리보기
            if doc.paragraphs:
                first_para = doc.paragraphs[0].text[:100] + "..." if len(doc.paragraphs[0].text) > 100 else doc.paragraphs[0].text
                print(f"✅ 첫 번째 문단 미리보기: {first_para}")
        else:
            print(f"❌ 테스트 파일이 없습니다: {test_file}")
            
    except ImportError as e:
        print(f"❌ python-docx 모듈 import 실패: {e}")
        return False
    except Exception as e:
        print(f"❌ DOCX 처리 오류: {e}")
        return False
    
    return True

def test_document_processor():
    """DocumentProcessor 테스트"""
    print("\n=== DocumentProcessor 테스트 ===")
    
    try:
        from services.document_processor import DocumentProcessor
        processor = DocumentProcessor()
        print("✅ DocumentProcessor 초기화 성공")
        
        # 지원되는 파일 형식 확인
        print(f"✅ 지원되는 파일 형식: {list(processor.loaders.keys())}")
        
        # DOCX 로더 확인
        if '.docx' in processor.loaders:
            print("✅ DOCX 로더 사용 가능")
            
            # 실제 파일로 테스트
            test_file = r"D:\99_DEOTIS_QA_SYSTEM\03_DEOTIS_QA\s3\BC카드(신용카드 업무처리 안내).docx"
            if os.path.exists(test_file):
                try:
                    documents = processor.load_document(test_file)
                    print(f"✅ 문서 로딩 성공 - {len(documents)}개 문서")
                    
                    if documents:
                        content_preview = documents[0].page_content[:200] + "..." if len(documents[0].page_content) > 200 else documents[0].page_content
                        print(f"✅ 내용 미리보기: {content_preview}")
                        
                        # 청킹 테스트
                        chunks = processor.split_documents(documents)
                        print(f"✅ 청킹 성공 - {len(chunks)}개 청크")
                        
                except Exception as e:
                    print(f"❌ 문서 처리 오류: {e}")
                    return False
            else:
                print("⚠️ 테스트 파일이 없어서 실제 로딩 테스트 생략")
        else:
            print("❌ DOCX 로더 사용 불가")
            return False
            
    except ImportError as e:
        print(f"❌ DocumentProcessor import 실패: {e}")
        return False
    except Exception as e:
        print(f"❌ DocumentProcessor 테스트 오류: {e}")
        return False
    
    return True

def test_full_loading():
    """전체 문서 로딩 프로세스 테스트"""
    print("\n=== 전체 문서 로딩 테스트 ===")
    
    try:
        from load_documents import load_s3_documents
        print("✅ load_s3_documents 함수 import 성공")
        
        # 실제 로딩 테스트
        documents_loaded, total_chunks = load_s3_documents()
        print(f"✅ 문서 로딩 완료")
        print(f"   - 로드된 문서: {documents_loaded}개")
        print(f"   - 생성된 청크: {total_chunks}개")
        
        if documents_loaded > 0 and total_chunks > 0:
            return True
        else:
            print("⚠️ 문서가 로드되지 않았습니다")
            return False
            
    except Exception as e:
        print(f"❌ 전체 로딩 테스트 오류: {e}")
        return False

if __name__ == "__main__":
    print("DOCX 문서 로딩 진단 시작...")
    print("=" * 50)
    
    all_passed = True
    
    # 1단계: DOCX 모듈 테스트
    if not test_docx_import():
        all_passed = False
    
    # 2단계: DocumentProcessor 테스트  
    if not test_document_processor():
        all_passed = False
    
    # 3단계: 전체 로딩 테스트
    if not test_full_loading():
        all_passed = False
    
    print("\n" + "=" * 50)
    if all_passed:
        print("🎉 모든 테스트 통과! 문서 로딩이 정상 작동합니다.")
        print("서버를 재시작하면 문서가 정상적으로 로드될 것입니다.")
    else:
        print("❌ 일부 테스트 실패. 위의 오류를 확인하고 해결하세요.")
    
    input("\n계속하려면 Enter를 누르세요...")