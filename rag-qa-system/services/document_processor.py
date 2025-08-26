from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import (
    TextLoader,
    UnstructuredMarkdownLoader
)
try:
    from langchain_community.document_loaders import PyPDFLoader
except ImportError:
    PyPDFLoader = None

# docx2txt 대신 python-docx 직접 사용
try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from langchain.schema import Document
from config import Config
import os
import time
from typing import List, Dict, Any, Callable
from services.chunking_strategies import get_chunking_strategy

class DocumentProcessor:
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=Config.CHUNK_SIZE,
            chunk_overlap=Config.CHUNK_OVERLAP,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
        self.loaders = {
            '.txt': TextLoader,
            '.md': UnstructuredMarkdownLoader
        }
        if PyPDFLoader:
            self.loaders['.pdf'] = PyPDFLoader
        if DOCX_AVAILABLE:
            self.loaders['.docx'] = self._load_docx_custom
    
    def _load_docx_custom(self, file_path: str) -> List[Document]:
        """python-docx를 사용한 DOCX 파일 로더"""
        try:
            doc = docx.Document(file_path)
            full_text = []
            
            # 문단 텍스트 추출
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text)
            
            # 표 텍스트 추출
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        full_text.append(" | ".join(row_text))
            
            content = "\n".join(full_text)
            
            if content.strip():
                metadata = {
                    'source': file_path,
                    'file_type': 'docx'
                }
                return [Document(page_content=content, metadata=metadata)]
            else:
                return []
                
        except Exception as e:
            raise Exception(f"Error loading DOCX document: {str(e)}")
    
    def load_document(self, file_path: str) -> List[Document]:
        """Load a document based on its file extension"""
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension not in self.loaders:
            raise ValueError(f"Unsupported file type: {file_extension}")
        
        loader_class_or_method = self.loaders[file_extension]
        
        try:
            if file_extension == '.docx':
                # 커스텀 DOCX 로더 사용
                documents = loader_class_or_method(file_path)
            else:
                # 기존 로더 사용
                loader = loader_class_or_method(file_path)
                documents = loader.load()
            return documents
        except Exception as e:
            raise Exception(f"Error loading document: {str(e)}")
    
    def split_documents(self, documents: List[Document]) -> List[Document]:
        """Split documents into chunks"""
        return self.text_splitter.split_documents(documents)
    
    def process_file(self, file_path: str, metadata: dict = None, chunking_strategy: str = "basic", 
                     progress_callback: Callable = None) -> Dict[str, Any]:
        """Process a file: load and split into chunks with progress tracking"""
        start_time = time.time()
        result = {
            "chunks": [],
            "chunking_strategy": chunking_strategy,
            "processing_time": 0,
            "file_size": 0,
            "chunk_count": 0,
            "error": None
        }
        
        try:
            if progress_callback:
                progress_callback("파일 로딩 중...", 10)
            
            # Load the document
            documents = self.load_document(file_path)
            result["file_size"] = os.path.getsize(file_path)
            
            if progress_callback:
                progress_callback("문서 로딩 완료, 청킹 전략 적용 중...", 30)
            
            # Add custom metadata if provided
            if metadata:
                for doc in documents:
                    doc.metadata.update(metadata)
            
            if progress_callback:
                progress_callback("청킹 처리 중...", 50)
            
            # Split into chunks using specified strategy
            chunks = self.split_documents_with_strategy(documents, chunking_strategy)
            
            if progress_callback:
                progress_callback("메타데이터 추가 중...", 80)
            
            # Add chunk index to metadata
            for i, chunk in enumerate(chunks):
                chunk.metadata['chunk_index'] = i
                chunk.metadata['source_file'] = os.path.basename(file_path)
                chunk.metadata['processing_time'] = time.strftime('%Y-%m-%d %H:%M:%S')
            
            result["chunks"] = chunks
            result["chunk_count"] = len(chunks)
            
            if progress_callback:
                progress_callback("처리 완료!", 100)
                
        except Exception as e:
            result["error"] = str(e)
            if progress_callback:
                progress_callback(f"오류: {str(e)}", 0)
        
        result["processing_time"] = time.time() - start_time
        return result
    
    def split_documents_with_strategy(self, documents: List[Document], strategy: str = "basic") -> List[Document]:
        """지정된 청킹 전략으로 문서 분할"""
        try:
            chunking_strategy = get_chunking_strategy(strategy)
            return chunking_strategy.split_documents(documents)
        except Exception as e:
            print(f"⚠️ 청킹 전략 '{strategy}' 실패, 기본 전략 사용: {e}")
            # 폴백: 기본 청킹 사용
            return self.text_splitter.split_documents(documents)
    
    def detect_chunking_strategy(self, text: str) -> str:
        """텍스트 내용을 분석하여 적절한 청킹 전략 감지"""
        # 커스텀 구분자 확인
        if "/$$/" in text:
            return "custom_delimiter"
        
        # 문서 크기에 따른 전략 선택
        text_length = len(text)
        if text_length > 10000:
            return "hybrid"  # 큰 문서는 하이브리드 전략
        elif text_length > 5000:
            return "semantic"  # 중간 크기는 의미 기반
        else:
            return "basic"  # 작은 문서는 기본 전략
    
    def process_text(self, text: str, metadata: dict = None) -> List[Document]:
        """Process raw text into chunks"""
        # Create a document from text
        doc = Document(page_content=text, metadata=metadata or {})
        
        # Split into chunks
        chunks = self.text_splitter.split_documents([doc])
        
        # Add chunk index to metadata
        for i, chunk in enumerate(chunks):
            chunk.metadata['chunk_index'] = i
        
        return chunks
    
    def validate_file(self, file_path: str) -> bool:
        """Validate if a file can be processed"""
        # Check if file exists
        if not os.path.exists(file_path):
            return False
        
        # Check file size
        file_size = os.path.getsize(file_path)
        if file_size > Config.MAX_FILE_SIZE:
            return False
        
        # Check file extension
        file_extension = os.path.splitext(file_path)[1].lower()
        if file_extension not in self.loaders:
            return False
        
        return True
