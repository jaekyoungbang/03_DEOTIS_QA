from typing import List, Dict, Any
import re
from langchain.text_splitter import RecursiveCharacterTextSplitter, CharacterTextSplitter
from langchain.schema import Document
import time
import os
from docx import Document as DocxDocument

class ChunkingBenchmarker:
    """청킹 전략 벤치마킹 시스템"""
    
    def __init__(self):
        self.strategies = {
            'basic': BasicChunkingStrategy(),
            'semantic': SemanticChunkingStrategy(),
            'hybrid': HybridChunkingStrategy()
        }
        self.results = []
    
    def benchmark_strategy(self, strategy_name: str, documents: List[Document]) -> Dict:
        """특정 청킹 전략 벤치마킹"""
        if strategy_name not in self.strategies:
            raise ValueError(f"Unknown strategy: {strategy_name}")
        
        strategy = self.strategies[strategy_name]
        start_time = time.time()
        
        # 메모리 사용량 측정 (간소화)
        start_memory = 0
        
        # 청킹 실행
        chunks = strategy.split_documents(documents)
        
        end_time = time.time()
        end_memory = 0
        
        # 품질 메트릭 계산
        quality_metrics = self._calculate_quality_metrics(chunks)
        
        result = {
            'strategy': strategy_name,
            'processing_time': end_time - start_time,
            'memory_used': end_memory - start_memory,
            'total_chunks': len(chunks),
            'avg_chunk_size': sum(len(chunk.page_content) for chunk in chunks) / len(chunks) if chunks else 0,
            'std_chunk_size': 0,  # 간소화
            'quality_metrics': quality_metrics,
            'timestamp': time.time()
        }
        
        self.results.append(result)
        return result
    
    def _calculate_quality_metrics(self, chunks: List[Document]) -> Dict:
        """청킹 품질 메트릭 계산"""
        if not chunks:
            return {'coherence_score': 0, 'completeness_score': 0}
        
        chunk_sizes = [len(chunk.page_content) for chunk in chunks]
        
        # 일관성 점수 (간소화된 버전)
        if chunk_sizes:
            avg_size = sum(chunk_sizes) / len(chunk_sizes)
            max_deviation = max(abs(size - avg_size) for size in chunk_sizes)
            coherence_score = max(0, 1 - (max_deviation / avg_size)) if avg_size > 0 else 1
        else:
            coherence_score = 0
        
        # 완성도 점수 (빈 청크나 너무 작은 청크가 적을수록 좋음)
        min_size_threshold = 100
        valid_chunks = sum(1 for size in chunk_sizes if size >= min_size_threshold)
        completeness_score = valid_chunks / len(chunks) if chunks else 0
        
        return {
            'coherence_score': coherence_score,
            'completeness_score': completeness_score,
            'avg_size': sum(chunk_sizes) / len(chunk_sizes) if chunk_sizes else 0,
            'size_variance': 0  # 간소화
        }
    
    def compare_strategies(self, documents: List[Document]) -> Dict:
        """모든 전략 비교"""
        comparison_results = {}
        
        for strategy_name in self.strategies.keys():
            try:
                result = self.benchmark_strategy(strategy_name, documents)
                comparison_results[strategy_name] = result
            except Exception as e:
                comparison_results[strategy_name] = {
                    'error': str(e),
                    'strategy': strategy_name
                }
        
        # 최적 전략 선택
        best_strategy = self._select_best_strategy(comparison_results)
        
        return {
            'results': comparison_results,
            'best_strategy': best_strategy,
            'comparison_summary': self._generate_comparison_summary(comparison_results)
        }
    
    def _select_best_strategy(self, results: Dict) -> str:
        """최적 전략 선택 (종합 점수 기반)"""
        scores = {}
        
        for strategy, result in results.items():
            if 'error' in result:
                scores[strategy] = 0
                continue
            
            # 종합 점수 계산 (속도 + 품질)
            speed_score = 1 / (result['processing_time'] + 0.1)  # 빠를수록 좋음
            quality_score = (
                result['quality_metrics']['coherence_score'] + 
                result['quality_metrics']['completeness_score']
            ) / 2
            
            scores[strategy] = (speed_score * 0.3) + (quality_score * 0.7)
        
        return max(scores, key=scores.get) if scores else 'basic'
    
    def _generate_comparison_summary(self, results: Dict) -> Dict:
        """비교 요약 생성"""
        summary = {
            'fastest': None,
            'highest_quality': None,
            'most_balanced': None
        }
        
        valid_results = {k: v for k, v in results.items() if 'error' not in v}
        
        if not valid_results:
            return summary
        
        # 가장 빠른 전략
        fastest = min(valid_results.items(), 
                     key=lambda x: x[1]['processing_time'])
        summary['fastest'] = fastest[0]
        
        # 가장 높은 품질
        highest_quality = max(valid_results.items(),
                            key=lambda x: (x[1]['quality_metrics']['coherence_score'] + 
                                         x[1]['quality_metrics']['completeness_score']) / 2)
        summary['highest_quality'] = highest_quality[0]
        
        # 가장 균형잡힌 전략 (이미 계산된 best_strategy와 동일)
        summary['most_balanced'] = self._select_best_strategy(results)
        
        return summary

class BasicChunkingStrategy:
    """기본 청킹 전략 - 문자 기반"""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        # Optimized for reduced LLM input length - smaller chunks work better
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        # 섹션 구분자를 추가하여 더 나은 청킹
        self.splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=[
                "\n\n\n",  # 섹션 간 구분
                "\n\n",    # 단락 구분
                "\n",      # 줄 구분
                ". ",      # 문장 구분
                " ",       # 단어 구분
                ""
            ]
        )
    
    def split_documents(self, documents: List[Document]) -> List[Document]:
        """문서를 청크로 분할"""
        # 섹션 기반 사전 처리
        enhanced_documents = self._preprocess_sections(documents)
        chunks = self.splitter.split_documents(enhanced_documents)
        
        # 제목 전용 청크 필터링 - 더 엄격한 기준 적용
        filtered_chunks = []
        for chunk in chunks:
            content = chunk.page_content.strip()
            
            # 제목 전용 청크인지 확인
            if self._is_title_only_chunk(content):
                print(f"   ❌ 제목 전용 청크 필터링: {repr(content[:50])}...")
                continue
                
            # 최소 길이 체크 (100자 이상)
            if len(content) >= 100:
                filtered_chunks.append(chunk)
            else:
                print(f"   ❌ 짧은 청크 필터링 ({len(content)}자): {repr(content[:30])}...")
        
        print(f"   📊 청킹 결과: {len(chunks)}개 → {len(filtered_chunks)}개 (필터링: {len(chunks) - len(filtered_chunks)}개)")
        return filtered_chunks if filtered_chunks else chunks  # 모든 청크가 필터링되면 원본 반환
    
    def _preprocess_sections(self, documents: List[Document]) -> List[Document]:
        """섹션 구분을 명확히 하기 위한 사전 처리"""
        processed_docs = []
        
        for doc in documents:
            content = doc.page_content
            
            # 숫자) 또는 숫자-숫자) 패턴으로 시작하는 섹션 구분
            # 예: "10) 회원제 업소", "1-1) 일시불 구매" 등
            section_pattern = r'(\n)(\d+[-]?\d*\))'
            content = re.sub(section_pattern, r'\n\n\n\2', content)
            
            # 대제목 구분 (예: [카드생활 TIP], /. 신용카드 TIP)
            # 제목이 단독으로 분리되지 않도록 구분자를 최소화
            title_pattern = r'(\n)(\[.*?\]|/\..*?)(\n)'
            # 제목 앞에만 구분을 두고, 뒤는 바로 이어지도록 함
            content = re.sub(title_pattern, r'\n\n\2\n', content)
            
            # A-1., A-2. 같은 섹션 구분
            section_letter_pattern = r'(\n)([A-Z]-\d+\.)'
            content = re.sub(section_letter_pattern, r'\n\n\n\2', content)
            
            processed_doc = Document(
                page_content=content,
                metadata=doc.metadata.copy()
            )
            processed_docs.append(processed_doc)
        
        return processed_docs
    
    def _is_title_only_chunk(self, content: str) -> bool:
        """제목 전용 청크인지 확인"""
        content = content.strip()
        
        # 정확히 "[민원접수방법 안내]" 같은 패턴만 있는 경우
        if content.startswith('[') and content.endswith(']') and '\n' not in content:
            return True
        
        # /. 로 시작하는 제목만 있는 경우  
        if content.startswith('/.') and len(content) < 100:
            return True
            
        # 줄바꿈이 없고 짧은 경우 (제목일 가능성)
        if '\n' not in content and len(content) < 50:
            return True
            
        # 매우 짧은 청크 (50자 미만)
        if len(content) < 50:
            return True
        
        # 제목으로만 구성된 경우 (예: "[민원접수방법 안내]"만 있고 다른 내용이 없는 경우)
        lines = [line.strip() for line in content.split('\n') if line.strip()]
        if len(lines) == 1 and lines[0].startswith('[') and lines[0].endswith(']'):
            return True
            
        return False

class SemanticChunkingStrategy:
    """의미 기반 청킹 전략 (간소화 버전)"""
    
    def __init__(self, min_chunk_size: int = 500, max_chunk_size: int = 1500):
        self.min_chunk_size = min_chunk_size
        self.max_chunk_size = max_chunk_size
        # 간소화된 버전 - 의존성 제거
        self.sentence_model = None
    
    def split_documents(self, documents: List[Document]) -> List[Document]:
        """의미 기반으로 문서 분할"""
        if not self.sentence_model:
            # 모델이 없으면 기본 전략 사용
            basic_strategy = BasicChunkingStrategy()
            return basic_strategy.split_documents(documents)
        
        all_chunks = []
        
        for doc in documents:
            chunks = self._split_document_semantically(doc)
            all_chunks.extend(chunks)
        
        return all_chunks
    
    def _split_document_semantically(self, document: Document) -> List[Document]:
        """단일 문서를 의미 기반으로 분할 (간소화)"""
        # 간소화된 버전 - 문장 구분자 기반 분할
        text = document.page_content
        
        # 간단한 문장 분할 (마침표 기준)
        sentences = [s.strip() for s in text.split('.') if s.strip()]
        
        if len(sentences) <= 1:
            return [document]
        
        # 크기 기반 그루핑
        chunks = self._group_sentences_by_size(sentences, document.metadata)
        
        return chunks
    
    def _group_sentences_by_size(self, sentences: List[str], metadata: Dict) -> List[Document]:
        """크기 기반 문장 그루핑 (간소화)"""
        chunks = []
        current_chunk = []
        current_size = 0
        
        for sentence in sentences:
            sentence_size = len(sentence)
            
            # 현재 청크에 추가할지 판단
            if current_size + sentence_size <= self.max_chunk_size:
                current_chunk.append(sentence + '.')  # 마침표 복원
                current_size += sentence_size
            else:
                # 현재 청크 완성
                if current_chunk and current_size >= self.min_chunk_size:
                    chunk_text = " ".join(current_chunk)
                    chunks.append(Document(
                        page_content=chunk_text,
                        metadata=metadata.copy()
                    ))
                
                # 새 청크 시작
                current_chunk = [sentence + '.']
                current_size = sentence_size
        
        # 마지막 청크 처리
        if current_chunk:
            chunk_text = " ".join(current_chunk)
            if len(chunk_text) >= self.min_chunk_size:
                chunks.append(Document(
                    page_content=chunk_text,
                    metadata=metadata.copy()
                ))
        
        return chunks if chunks else [Document(
            page_content=" ".join(s + '.' for s in sentences),
            metadata=metadata.copy()
        )]

class HybridChunkingStrategy:
    """하이브리드 청킹 전략 - 기본 + 의미 기반"""
    
    def __init__(self):
        self.basic_strategy = BasicChunkingStrategy(chunk_size=800, chunk_overlap=150)
        self.semantic_strategy = SemanticChunkingStrategy(min_chunk_size=400, max_chunk_size=1200)
    
    def split_documents(self, documents: List[Document]) -> List[Document]:
        """하이브리드 방식으로 문서 분할"""
        # 먼저 기본 전략으로 분할
        basic_chunks = self.basic_strategy.split_documents(documents)
        
        # 큰 청크들은 의미 기반으로 재분할
        final_chunks = []
        
        for chunk in basic_chunks:
            if len(chunk.page_content) > 1200:
                # 큰 청크는 의미 기반으로 재분할
                semantic_chunks = self.semantic_strategy._split_document_semantically(chunk)
                final_chunks.extend(semantic_chunks)
            else:
                final_chunks.append(chunk)
        
        return final_chunks

class S3BasicChunkingStrategy:
    """기본 청킹 전략 (s3 폴더 기반)"""
    
    def __init__(self):
        self.s3_docs_path = '/mnt/d/99_DEOTIS_QA_SYSTEM/03_DEOTIS_QA/s3/'
        self.basic_splitter = RecursiveCharacterTextSplitter(
            chunk_size=800,  # Reduced for LLM optimization
            chunk_overlap=150,  # Proportionally reduced overlap
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
    
    def split_documents(self, documents: List[Document]) -> List[Document]:
        """기본 청킹 수행"""
        basic_chunks = []
        
        for document in documents:
            # 기본 청킹 수행
            chunks = self.basic_splitter.split_documents([document])
            
            for chunk in chunks:
                chunk.metadata['chunking_strategy'] = 's3-basic'
                chunk.metadata['data_source'] = 's3'
                basic_chunks.append(chunk)
        
        return basic_chunks

class S3CustomChunkingStrategy:
    """커스텀 청킹 전략 (s3-chunking 폴더 기반)"""
    
    def __init__(self):
        self.s3_chunking_path = '/mnt/d/99_DEOTIS_QA_SYSTEM/03_DEOTIS_QA/s3-chunking/'
        self.card_knowledge = self._load_card_knowledge()
        self.basic_splitter = RecursiveCharacterTextSplitter(
            chunk_size=800,  # Reduced for LLM optimization
            chunk_overlap=150,  # Proportionally reduced overlap
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
    
    def _load_card_knowledge(self) -> Dict[str, str]:
        """전랩 카드 문서들에서 지식 추출 (s3-chunking 폴더)"""
        knowledge = {}
        
        try:
            # BC카드 신용카드 업무처리 안내
            doc1_path = os.path.join(self.s3_chunking_path, 'BC카드(신용카드 업무처리 안내).docx')
            if os.path.exists(doc1_path):
                doc1 = DocxDocument(doc1_path)
                content1 = '\n'.join([paragraph.text for paragraph in doc1.paragraphs if paragraph.text.strip()])
                knowledge['business_guide'] = content1
                print(f"[OK] s3-chunking 업무처리 안내 문서 로드: {len(content1)}자")
            
            # BC카드 카드이용안내
            doc2_path = os.path.join(self.s3_chunking_path, 'BC카드(카드이용안내).docx')
            if os.path.exists(doc2_path):
                doc2 = DocxDocument(doc2_path)
                content2 = '\n'.join([paragraph.text for paragraph in doc2.paragraphs if paragraph.text.strip()])
                knowledge['usage_guide'] = content2
                print(f"[OK] s3-chunking 이용안내 문서 로드: {len(content2)}자")
                
        except Exception as e:
            print(f"[WARNING] s3-chunking 전랩 문서 로드 오류: {e}")
        
        return knowledge
    
    def split_documents(self, documents: List[Document]) -> List[Document]:
        """전랩 지식을 활용한 커스텀 청킹"""
        enhanced_chunks = []
        
        for document in documents:
            # 기본 청킹 수행
            basic_chunks = self.basic_splitter.split_documents([document])
            
            # 각 청크를 전랩 지식으로 향상
            for chunk in basic_chunks:
                enhanced_chunk = self._enhance_chunk_with_s3_knowledge(chunk)
                enhanced_chunks.append(enhanced_chunk)
        
        return enhanced_chunks
    
    def _enhance_chunk_with_s3_knowledge(self, chunk: Document) -> Document:
        """청크를 전랩 지식으로 향상"""
        content = chunk.page_content
        metadata = chunk.metadata.copy()
        
        # 카드 관련 키워드 감지
        card_keywords = [
            '카드', '신용카드', '체크카드', '전용카드',
            '비밀번호', '이용한도', '수수료', '이자',
            '연체료', '결제', '승인', '취소'
        ]
        
        matched_keywords = [kw for kw in card_keywords if kw in content]
        
        if matched_keywords:
            # 전랩 배경 지식 추가
            enhanced_content = content
            
            # 업무처리 관련 지식 추가
            if any(kw in content for kw in ['승인', '취소', '비밀번호']):
                if 'business_guide' in self.card_knowledge:
                    relevant_info = self._extract_relevant_info(
                        content, self.card_knowledge['business_guide']
                    )
                    if relevant_info:
                        enhanced_content += f"\n\n[전랥 업무처리 지식]: {relevant_info[:500]}..."
            
            # 이용안내 관련 지식 추가
            if any(kw in content for kw in ['이용한도', '수수료', '이자']):
                if 'usage_guide' in self.card_knowledge:
                    relevant_info = self._extract_relevant_info(
                        content, self.card_knowledge['usage_guide']
                    )
                    if relevant_info:
                        enhanced_content += f"\n\n[전랥 이용안내 지식]: {relevant_info[:500]}..."
            
            metadata['enhanced_with_s3'] = True
            metadata['matched_keywords'] = matched_keywords
            metadata['chunking_strategy'] = 's3-custom'
            metadata['data_source'] = 's3-chunking'
            
            return Document(
                page_content=enhanced_content,
                metadata=metadata
            )
        
        # 키워드 매치가 없으면 기본 청크 반환
        metadata['chunking_strategy'] = 's3-custom-basic'
        metadata['data_source'] = 's3-chunking'
        return Document(
            page_content=content,
            metadata=metadata
        )
    
    def _extract_relevant_info(self, query_content: str, knowledge_base: str) -> str:
        """쿼리 내용과 관련된 지식 추출"""
        # 간단한 키워드 매칭 기반 추출
        query_words = re.findall(r'\w+', query_content.lower())
        knowledge_sentences = knowledge_base.split('\n')
        
        relevant_sentences = []
        for sentence in knowledge_sentences:
            if len(sentence.strip()) > 20:  # 충분히 긴 문장만
                sentence_lower = sentence.lower()
                matches = sum(1 for word in query_words if word in sentence_lower)
                if matches >= 2:  # 2개 이상 단어 매칭
                    relevant_sentences.append(sentence.strip())
        
        return '. '.join(relevant_sentences[:3])  # 상위 3개 문장

# 전역 벤치마커 인스턴스
chunking_benchmarker = ChunkingBenchmarker()

class CustomDelimiterChunkingStrategy:
    """커스텀 구분자 청킹 전략 - /$$/ 구분자 기반"""
    
    def __init__(self, delimiter: str = "/$$/" , fallback_chunk_size: int = 1000, fallback_overlap: int = 200):
        self.delimiter = delimiter
        self.fallback_splitter = RecursiveCharacterTextSplitter(
            chunk_size=fallback_chunk_size,
            chunk_overlap=fallback_overlap,
            separators=["\n\n", "\n", " ", ""]
        )
    
    def split_documents(self, documents: List[Document]) -> List[Document]:
        """커스텀 구분자로 문서 분할"""
        all_chunks = []
        
        for doc in documents:
            chunks = self._split_document_by_delimiter(doc)
            all_chunks.extend(chunks)
        
        return all_chunks
    
    def _split_document_by_delimiter(self, document: Document) -> List[Document]:
        """단일 문서를 커스텀 구분자로 분할"""
        text = document.page_content
        metadata = document.metadata.copy()
        
        # 커스텀 구분자가 있는지 확인
        if self.delimiter in text:
            # 커스텀 구분자로 분할
            parts = text.split(self.delimiter)
            chunks = []
            
            for i, part in enumerate(parts):
                part = part.strip()
                if part:  # 빈 부분 제외
                    chunk_metadata = metadata.copy()
                    chunk_metadata['chunking_strategy'] = 'custom_delimiter'
                    chunk_metadata['chunk_index'] = i
                    chunk_metadata['total_chunks'] = len([p for p in parts if p.strip()])
                    chunk_metadata['delimiter_used'] = self.delimiter
                    
                    chunks.append(Document(
                        page_content=part,
                        metadata=chunk_metadata
                    ))
            
            return chunks if chunks else self._fallback_chunking(document)
        else:
            # 구분자가 없으면 기본 청킹 사용
            return self._fallback_chunking(document)
    
    def _fallback_chunking(self, document: Document) -> List[Document]:
        """폴백 청킹 (기본 전략)"""
        chunks = self.fallback_splitter.split_documents([document])
        
        for chunk in chunks:
            chunk.metadata['chunking_strategy'] = 'basic_fallback'
            chunk.metadata['delimiter_used'] = 'none'
        
        return chunks

def get_chunking_strategy(strategy_name: str = 'basic'):
    """청킹 전략 팩토리 함수"""
    strategies = {
        'basic': BasicChunkingStrategy(),
        'semantic': SemanticChunkingStrategy(),
        'hybrid': HybridChunkingStrategy(),
        's3-basic': S3BasicChunkingStrategy(),
        's3-custom': S3CustomChunkingStrategy(),
        'custom_delimiter': CustomDelimiterChunkingStrategy()
    }
    
    return strategies.get(strategy_name, strategies['basic'])

def benchmark_chunking_strategies(documents: List[Document]) -> Dict:
    """청킹 전략 벤치마킹 수행"""
    return chunking_benchmarker.compare_strategies(documents)