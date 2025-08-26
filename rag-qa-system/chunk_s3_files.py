#!/usr/bin/env python3
"""
S3-CHUNKING 폴더의 파일들을 /$$/ 구분자로 청킹하는 스크립트
"""

import os
from docx import Document
import json

def chunk_docx_file(file_path, delimiter="/$$/"):
    """DOCX 파일을 구분자 기준으로 청킹"""
    try:
        doc = Document(file_path)
        
        # 전체 텍스트 추출
        full_text = []
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                full_text.append(text)
        
        # 텍스트 결합
        combined_text = '\n'.join(full_text)
        
        # 구분자로 청킹
        if delimiter in combined_text:
            chunks = combined_text.split(delimiter)
            print(f"✅ {os.path.basename(file_path)}: {len(chunks)}개 청크로 분할됨 (구분자: {delimiter})")
            return chunks
        else:
            print(f"⚠️  {os.path.basename(file_path)}: 구분자 '{delimiter}'를 찾을 수 없음")
            # 구분자가 없으면 전체를 하나의 청크로
            return [combined_text]
            
    except Exception as e:
        print(f"❌ {file_path} 처리 중 오류: {str(e)}")
        return []

def add_delimiter_to_docx(file_path, output_path, sections, delimiter="/$$/"):
    """DOCX 파일에 구분자를 추가하여 저장"""
    try:
        doc = Document(file_path)
        new_doc = Document()
        
        # 문서를 섹션별로 나누어 구분자 추가
        current_section = []
        section_keywords = sections  # 섹션을 나누는 키워드들
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            
            # 섹션 키워드를 찾으면
            is_section_start = False
            for keyword in section_keywords:
                if keyword in text:
                    # 이전 섹션 저장
                    if current_section:
                        for para_text in current_section:
                            new_doc.add_paragraph(para_text)
                        # 구분자 추가
                        new_doc.add_paragraph(delimiter)
                        current_section = []
                    is_section_start = True
                    break
            
            # 현재 단락 추가
            if text:
                current_section.append(text)
        
        # 마지막 섹션 저장
        if current_section:
            for para_text in current_section:
                new_doc.add_paragraph(para_text)
        
        # 파일 저장
        new_doc.save(output_path)
        print(f"✅ 구분자가 추가된 파일 저장됨: {output_path}")
        return True
        
    except Exception as e:
        print(f"❌ 구분자 추가 중 오류: {str(e)}")
        return False

def process_s3_chunking_folder():
    """s3-chunking 폴더의 모든 파일 처리"""
    s3_folder = "/mnt/d/99_DEOTIS_QA_SYSTEM/03_DEOTIS_QA/s3-chunking"
    output_folder = "/mnt/d/99_DEOTIS_QA_SYSTEM/03_DEOTIS_QA/rag-qa-system/data/chunked_documents"
    
    # 출력 폴더 생성
    os.makedirs(output_folder, exist_ok=True)
    
    print("="*60)
    print("S3-CHUNKING 폴더 파일 청킹 시작")
    print("="*60)
    
    # BC카드 문서를 위한 섹션 키워드 정의
    bc_card_sections = [
        "1. 개요",
        "2. 신용카드 발급",
        "3. 카드 이용",
        "4. 결제",
        "5. 부가서비스",
        "6. 고객서비스",
        "제1장",
        "제2장",
        "제3장",
        "제4장",
        "제5장",
        "■",  # 주요 섹션 마커
        "▶",  # 서브 섹션 마커
    ]
    
    # 폴더 내 모든 파일 처리
    for filename in os.listdir(s3_folder):
        if filename.endswith('.docx'):
            file_path = os.path.join(s3_folder, filename)
            
            print(f"\n📄 처리 중: {filename}")
            
            # 1. 원본 파일 청킹 시도
            chunks = chunk_docx_file(file_path)
            
            # 2. 구분자가 없으면 추가
            if len(chunks) <= 1:
                print(f"💡 구분자를 추가하여 재청킹 시도...")
                
                # 구분자가 추가된 새 파일 생성
                modified_filename = f"chunked_{filename}"
                modified_path = os.path.join(output_folder, modified_filename)
                
                if add_delimiter_to_docx(file_path, modified_path, bc_card_sections):
                    # 수정된 파일 다시 청킹
                    chunks = chunk_docx_file(modified_path)
            
            # 3. 청킹 결과 저장
            if chunks and len(chunks) > 1:
                result_filename = f"{filename.replace('.docx', '')}_chunks.json"
                result_path = os.path.join(output_folder, result_filename)
                
                chunk_data = {
                    'source_file': filename,
                    'delimiter': "/$$/",
                    'chunk_count': len(chunks),
                    'chunks': []
                }
                
                for i, chunk in enumerate(chunks):
                    chunk_info = {
                        'index': i,
                        'content': chunk.strip(),
                        'length': len(chunk.strip())
                    }
                    chunk_data['chunks'].append(chunk_info)
                
                # JSON으로 저장
                with open(result_path, 'w', encoding='utf-8') as f:
                    json.dump(chunk_data, f, ensure_ascii=False, indent=2)
                
                print(f"✅ 청킹 결과 저장: {result_filename}")
                
                # 청크 요약 출력
                print(f"   - 총 청크 수: {len(chunks)}")
                for i, chunk in enumerate(chunks[:3]):  # 처음 3개만 미리보기
                    preview = chunk.strip()[:100] + "..." if len(chunk.strip()) > 100 else chunk.strip()
                    print(f"   - 청크 {i+1}: {preview}")
                
                if len(chunks) > 3:
                    print(f"   ... 외 {len(chunks)-3}개 청크")

if __name__ == "__main__":
    process_s3_chunking_folder()