#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word 파일에 /$$/ 구분자를 추가하는 도구
"""

import os
import sys
import re
from docx import Document
from pathlib import Path

def add_delimiters_to_word_file(input_path, output_path=None):
    """Word 파일에 /$$/ 구분자를 추가"""
    
    if output_path is None:
        # 원본 파일명에 _delimited 추가
        path = Path(input_path)
        output_path = str(path.parent / f"{path.stem}_delimited{path.suffix}")
    
    print(f"📄 파일 처리 중: {input_path}")
    
    # Word 문서 열기
    doc = Document(input_path)
    
    # 새 문서 생성
    new_doc = Document()
    
    processed_paragraphs = 0
    added_delimiters = 0
    
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        
        # 빈 단락 건너뛰기
        if not text:
            continue
        
        # 구분이 필요한 패턴들 정의
        section_patterns = [
            r'^\d+[-]?\d*\)',  # 숫자) 또는 숫자-숫자) 패턴
            r'^[A-Z]-\d+\.',   # A-1., A-2. 같은 패턴
            r'^\[.*?\]',       # [제목] 패턴
            r'^/\..*',         # /. 로 시작하는 패턴
        ]
        
        # 섹션 시작인지 확인
        is_section_start = any(re.match(pattern, text) for pattern in section_patterns)
        
        # 이전 단락이 있고 현재가 섹션 시작이면 구분자 추가
        if processed_paragraphs > 0 and is_section_start:
            delimiter_paragraph = new_doc.add_paragraph("/$$/ ")
            added_delimiters += 1
            print(f"   구분자 추가: {text[:50]}...")
        
        # 현재 단락 추가
        new_paragraph = new_doc.add_paragraph(text)
        processed_paragraphs += 1
    
    # 문서 저장
    new_doc.save(output_path)
    
    print(f"✅ 처리 완료:")
    print(f"   - 입력 파일: {input_path}")
    print(f"   - 출력 파일: {output_path}")
    print(f"   - 처리된 단락: {processed_paragraphs}개")
    print(f"   - 추가된 구분자: {added_delimiters}개")
    
    return output_path

def process_s3_chunking_folder():
    """s3-chunking 폴더의 모든 Word 파일 처리"""
    
    s3_chunking_path = "/mnt/d/99_DEOTIS_QA_SYSTEM/03_DEOTIS_QA/s3-chunking"
    
    print(f"📂 s3-chunking 폴더 처리: {s3_chunking_path}")
    
    if not os.path.exists(s3_chunking_path):
        print(f"❌ 폴더가 존재하지 않습니다: {s3_chunking_path}")
        return
    
    processed_files = []
    
    # Word 파일 찾기
    for file in os.listdir(s3_chunking_path):
        if file.endswith('.docx') and not file.startswith('~$'):
            input_path = os.path.join(s3_chunking_path, file)
            
            # 이미 delimited 버전이 있는지 확인
            file_stem = Path(file).stem
            if '_delimited' in file_stem:
                print(f"⏭️ 이미 처리된 파일 건너뛰기: {file}")
                continue
            
            try:
                output_path = add_delimiters_to_word_file(input_path)
                processed_files.append(output_path)
            except Exception as e:
                print(f"❌ 파일 처리 실패 ({file}): {e}")
    
    print(f"\n📊 전체 처리 완료:")
    print(f"   - 처리된 파일 수: {len(processed_files)}개")
    for file_path in processed_files:
        print(f"   - {os.path.basename(file_path)}")

def create_sample_delimited_file():
    """샘플로 구분자가 포함된 파일 생성"""
    
    sample_path = "/mnt/d/99_DEOTIS_QA_SYSTEM/03_DEOTIS_QA/s3-chunking/sample_with_delimiters.docx"
    
    print(f"📝 샘플 파일 생성: {sample_path}")
    
    # 새 문서 생성
    doc = Document()
    
    # 샘플 내용 추가
    doc.add_paragraph("신용카드 업무처리 안내")
    doc.add_paragraph("/$$/ ")
    
    doc.add_paragraph("A-1. [카드발급절차]")
    doc.add_paragraph("카드발급 신청시 카드발급 상담사는 발급불가조건만을 조회한 후 '카드발급이 가능합니다.'라고 안내합니다.")
    doc.add_paragraph("/$$/ ")
    
    doc.add_paragraph("A-2. [카드발급기준]")
    doc.add_paragraph("카드발급업무는 신용대출심사의 성격을 가지므로 카드회사는 카드발급 여부를 자율적으로 결정할 수 있습니다.")
    doc.add_paragraph("/$$/ ")
    
    doc.add_paragraph("1) 신용카드알뜰이용법")
    doc.add_paragraph("1-1) 일시불 구매")
    doc.add_paragraph("큰액수는 결제 시에 가급적 일시불로 이용하여 주시기 바랍니다.")
    doc.add_paragraph("/$$/ ")
    
    doc.add_paragraph("10) 회원제 업소 등의 경우 신용카드 할부 이용")
    doc.add_paragraph("회원제 업소 등의 경우 신용카드 할부 이용은 할부거래법 등에 의해 카드사가 보호하고 있습니다.")
    
    # 저장
    doc.save(sample_path)
    print(f"✅ 샘플 파일 생성 완료: {sample_path}")
    
    return sample_path

if __name__ == "__main__":
    print("=== Word 파일 /$$/ 구분자 추가 도구 ===\n")
    
    # 1. 샘플 파일 생성
    print("1. 샘플 파일 생성:")
    create_sample_delimited_file()
    
    print("\n" + "="*50 + "\n")
    
    # 2. 기존 파일들 처리
    print("2. 기존 Word 파일들 처리:")
    process_s3_chunking_folder()