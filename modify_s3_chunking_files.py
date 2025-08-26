#!/usr/bin/env python3
"""
s3-chunking 폴더의 Word 파일들에 /$/ 구분자를 직접 추가하는 스크립트
"""

import os
from docx import Document
import shutil

def add_delimiters_to_word_file(file_path, output_path, delimiter="/$$/"):
    """Word 파일에 구분자를 추가하여 새로운 파일로 저장"""
    try:
        # 원본 파일 읽기
        doc = Document(file_path)
        
        # 새 문서 생성
        new_doc = Document()
        
        # BC카드 문서를 위한 섹션 키워드 정의
        section_keywords = [
            "1. 개요",
            "2. 신용카드 발급",
            "3. 카드 이용",
            "4. 결제",
            "5. 부가서비스",
            "6. 고객서비스",
            "제1장",
            "제2장", 
            "제3장",
            "제4장",
            "제5장",
            "■",  # 주요 섹션 마커
            "▶",  # 서브 섹션 마커
        ]
        
        current_section = []
        section_count = 0
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            
            # 섹션 시작 키워드 확인
            is_section_start = False
            for keyword in section_keywords:
                if keyword in text and len(text) < 100:  # 제목은 보통 짧음
                    # 이전 섹션이 있으면 저장
                    if current_section:
                        for para_text in current_section:
                            if para_text.strip():
                                new_doc.add_paragraph(para_text)
                        
                        # 구분자 추가 (첫 번째 섹션이 아닌 경우)
                        if section_count > 0:
                            new_doc.add_paragraph(delimiter)
                            print(f"  └─ 구분자 추가됨: 섹션 {section_count}")
                        
                        section_count += 1
                        current_section = []
                    
                    is_section_start = True
                    break
            
            # 현재 단락을 섹션에 추가
            if text:
                current_section.append(text)
        
        # 마지막 섹션 저장
        if current_section:
            for para_text in current_section:
                if para_text.strip():
                    new_doc.add_paragraph(para_text)
            section_count += 1
        
        # 파일 저장
        new_doc.save(output_path)
        print(f"✅ 구분자가 추가된 파일 저장: {os.path.basename(output_path)}")
        print(f"   └─ 총 {section_count}개 섹션 생성")
        
        return True
        
    except Exception as e:
        print(f"❌ {file_path} 처리 중 오류: {str(e)}")
        return False

def backup_original_files(s3_folder):
    """원본 파일 백업"""
    backup_folder = os.path.join(s3_folder, "backup")
    os.makedirs(backup_folder, exist_ok=True)
    
    print("\n📦 원본 파일 백업 중...")
    
    for filename in os.listdir(s3_folder):
        if filename.endswith('.docx'):
            src_path = os.path.join(s3_folder, filename)
            backup_path = os.path.join(backup_folder, f"original_{filename}")
            
            if not os.path.exists(backup_path):
                shutil.copy2(src_path, backup_path)
                print(f"  ✅ 백업됨: {filename} → original_{filename}")

def modify_s3_chunking_files():
    """s3-chunking 폴더의 모든 Word 파일 수정"""
    s3_folder = "/mnt/d/99_DEOTIS_QA_SYSTEM/03_DEOTIS_QA/s3-chunking"
    
    print("=" * 60)
    print("s3-chunking 폴더 Word 파일에 /$/ 구분자 추가")
    print("=" * 60)
    
    # 1. 원본 파일 백업
    backup_original_files(s3_folder)
    
    print(f"\n📁 처리 대상 폴더: {s3_folder}")
    
    # 2. 각 Word 파일 처리
    for filename in os.listdir(s3_folder):
        if filename.endswith('.docx') and not filename.startswith('original_'):
            file_path = os.path.join(s3_folder, filename)
            
            print(f"\n📄 처리 중: {filename}")
            
            # 임시 파일로 구분자 추가된 버전 생성
            temp_path = os.path.join(s3_folder, f"temp_{filename}")
            
            if add_delimiters_to_word_file(file_path, temp_path):
                # 성공하면 원본을 대체
                try:
                    os.replace(temp_path, file_path)
                    print(f"  ✅ 원본 파일 업데이트 완료: {filename}")
                except Exception as e:
                    print(f"  ❌ 파일 교체 실패: {e}")
                    # 임시 파일 정리
                    if os.path.exists(temp_path):
                        os.remove(temp_path)
            else:
                # 실패하면 임시 파일 정리
                if os.path.exists(temp_path):
                    os.remove(temp_path)
    
    print(f"\n✅ s3-chunking 폴더의 모든 Word 파일이 /$/ 구분자로 수정되었습니다.")
    print(f"📦 원본 파일들은 backup 폴더에 보관되어 있습니다.")

if __name__ == "__main__":
    modify_s3_chunking_files()