#!/usr/bin/env python3
"""
새 파일명으로 구분자 추가된 파일 생성
"""

import os
from docx import Document

def create_chunked_file():
    """새로운 파일명으로 구분자가 추가된 파일 생성"""
    s3_folder = "/mnt/d/99_DEOTIS_QA_SYSTEM/03_DEOTIS_QA/s3-chunking"
    backup_folder = os.path.join(s3_folder, "backup")
    
    # 원본에서 새 파일 생성
    original_file = os.path.join(backup_folder, "original_BC카드(카드이용안내).docx")
    new_filename = "BC카드(카드이용안내)_chunked.docx"
    new_file_path = os.path.join(s3_folder, new_filename)
    
    print(f"📄 새 파일 생성: {new_filename}")
    
    try:
        # 1. 원본 문서 읽기
        doc = Document(original_file)
        
        # 2. 텍스트 추출
        all_text = []
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                all_text.append(text)
        
        print(f"  └─ 총 {len(all_text)}개 문단")
        
        # 3. 새 문서에 구분자 추가하여 저장
        new_doc = Document()
        
        # 전체를 8개 청크로 나누기
        chunk_size = len(all_text) // 8
        
        for i, text in enumerate(all_text):
            new_doc.add_paragraph(text)
            
            # 매 chunk_size 마다 구분자 추가
            if (i + 1) % chunk_size == 0 and i < len(all_text) - chunk_size:
                new_doc.add_paragraph("/$$/")
                print(f"    ├─ 구분자 추가: {(i+1)//chunk_size}번째 위치")
        
        # 4. 새 파일로 저장
        new_doc.save(new_file_path)
        print(f"✅ 새 파일 생성 완료: {new_filename}")
        print(f"   └─ 8개 청크로 분할됨")
        
        return True
        
    except Exception as e:
        print(f"❌ 오류: {str(e)}")
        return False

if __name__ == "__main__":
    create_chunked_file()